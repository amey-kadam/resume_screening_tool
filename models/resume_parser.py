import json
import os
import logging
from logging.handlers import RotatingFileHandler
import google.generativeai as genai
from google.generativeai import GenerativeModel
import docx
import PyPDF2
import fitz  
from io import BytesIO


# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
handler = RotatingFileHandler('resume_parser.log', maxBytes=10000, backupCount=1)
handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

# Initialize Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
logger.info("Google API configured")

model = GenerativeModel('gemini-1.5-flash-8b-001')
logger.info("Gemini model initialized")


async def extract_text_with_gemini(file_path):
    logger.info(f"Extracting text from {file_path}")

    try:
        with open(file_path, 'rb') as file:
            file_stream = BytesIO(file.read())

        if file_path.endswith('.pdf'):
            doc = fitz.open("pdf", file_stream.read())
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            logger.info("PDF file processed")
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_stream)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            logger.info("DOCX file processed")
        else:
            raise ValueError("Unsupported file type")

        logger.info("Text extraction completed")
        return full_text

    except Exception as e:
        logger.error(f"Error during text extraction: {str(e)}")
        raise

def parse_resume(file_path):
    logger.info(f"Parsing resume: {file_path}")
    extracted_text = extract_text_with_gemini(file_path)

    if not extracted_text.strip():
        logger.error("Extracted text is empty. Cannot parse.")
        return {
            "Name": "",
            "Skills": [],
            "Education": [],
            "Projects": [],
            "Experience": []
        }
    
    prompt = f"""
    Parse the following resume text and extract the relevant information as a JSON object with these keys:
    
    - "Name": Full name of the individual.
    - "Skills": A list of skills or expertise mentioned.
    - "Education": A list of educational qualifications (degree, institution, graduation year).
    - "Projects": A list of projects (title, description, technologies used, duration).
    - "Experience": A list of work experiences (job title, company, duration, responsibilities).
    
    Ensure the output is in valid JSON format. If a section is not found, return an empty list for that key.
    
    Resume Text:
    {extracted_text}
    
    Return only the JSON object, nothing else:
    """

    response = model.generate_content(prompt)
    logger.info("Resume parsing completed")

    try:
        parsed_data = json.loads(response.text)
        logger.info("JSON parsed successfully")
    except json.JSONDecodeError:
        logger.error("Error: Unable to parse JSON from Gemini response")
        logger.error(f"Gemini response: {response.text}")
        parsed_data = {
            "Name": "",
            "Skills": [],
            "Education": [],
            "Projects": [],
            "Experience": []
        }

    return parsed_data

def save_to_json(data, filename='resumes.json'):
    file_path = os.path.join(os.getcwd(), filename)
    logger.info(f"Saving data to {file_path}")

    try:
        # Load existing data from file, or create an empty list if file not found
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                existing_data = json.load(f)
            logger.info("Existing data loaded from file")
        else:
            existing_data = []
            logger.info("No existing file found, creating new data")

        # Optional: Check for duplicates (based on a unique identifier like Name or Email)
        if data not in existing_data:
            existing_data.append(data)
            logger.info("New data appended")
        else:
            logger.info("Duplicate entry detected. Not adding to the file.")
            return

        # Write updated data back to the file
        with open(file_path, 'w') as f:
            json.dump(existing_data, f, indent=2)

        logger.info(f"Data successfully saved to {file_path}")

    except json.JSONDecodeError as e:
        logger.error(f"Error reading JSON data: {e}")
    except IOError as e:
        logger.error(f"File I/O error: {e}")
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")

def process_resume(file_path):
    logger.info(f"Processing resume: {file_path}")
    parsed_data = parse_resume(file_path)
    save_to_json(parsed_data)
    logger.info("Resume processing completed")
    return parsed_data

def fallback_process_resume(file_path):
    logger.info(f"Using fallback process for resume: {file_path}")
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        logger.info("DOCX file processed")
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            full_text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
        logger.info("PDF file processed")
    else:
        logger.error("Unsupported file type")
        raise ValueError("Unsupported file type")

    name = full_text.split('\n')[0]  # Example of extracting name from text
    skills = ['Python', 'Flask']  # Fallback skills

    logger.info("Fallback processing completed")
    return {
        'Name': name,
        'Skills': skills,
    }
